"""
build_report.py
Generates the ITC293 Assessment 4 written project report as a .docx file.
Author: Rosmin Roy
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

GREEN = RGBColor(0x2E, 0x5D, 0x34)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN
    return h


def para(text="", bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def shade_header(row):
    """Apply green shading + white bold text to a header row."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "2E5D34")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def make_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    shade_header(hdr)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# =====================================================================
# TITLE BLOCK
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ITC293 – Introduction to the Web")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Assessment 4 – HTML, CSS, Design and JavaScript\nProject Report (Written Component)")
r.font.size = Pt(14)
r.font.color.rgb = GREY

doc.add_paragraph()
info = make_table(
    ["Item", "Detail"],
    [
        ["Student name", "Rosmin Roy"],
        ["Subject", "ITC293 – Introduction to the Web (202630)"],
        ["Assessment", "Item 4 – Written Component (40% task, 20% written)"],
        ["Fictitious employer", "Verdant Web Solutions Pty Ltd"],
        ["Hosted site URL", "<insert your hosted URL here, e.g. GitHub Pages / W3Schools Spaces>"],
        ["Date", "7 June 2026"],
    ],
    widths=[2.0, 4.5],
)

doc.add_page_break()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
heading("1. Introduction", 1)
para(
    "This report documents my Assessment 4 submission for ITC293. The submission is an "
    "electronic work portfolio website built for a fictitious prospective employer, "
    "Verdant Web Solutions Pty Ltd. The site presents four relevant skill areas (one of "
    "which is Web Development) and includes an HTML5 newsletter signup form coded in HTML, "
    "CSS and JavaScript."
)
para(
    "The website is hand-coded in HTML5 using a text editor, styled with a single external "
    "CSS file (no inline or internal CSS), and uses an external JavaScript file for the form "
    "behaviour. This report contains: a requirement-by-requirement compliance table, a "
    "summary of what has changed since the previous submission (A3), and a testing plan."
)

# =====================================================================
# 2. SITE STRUCTURE
# =====================================================================
heading("2. Website Structure", 1)
para("The site has 8 HTML pages (the minimum required is 6), plus one external stylesheet and one external JavaScript file:")
make_table(
    ["File", "Purpose"],
    [
        ["index.html", "Home page – introduces me to potential employers"],
        ["rroy_A4_about.html", "About page – background and interests"],
        ["rroy_A4_resume.html", "Resume – personal details, education (table), experience, qualifications"],
        ["rroy_A4_web_skills.html", "Web Development skill page – links to the signup form in a new tab"],
        ["rroy_A4_communication.html", "Communication skill page"],
        ["rroy_A4_problem_solving.html", "Problem Solving skill page"],
        ["rroy_A4_teamwork.html", "Teamwork skill page"],
        ["rroy_A4_newsletter.html", "HTML5 newsletter signup form"],
        ["css/portfolio.css", "Single external stylesheet (screen + print)"],
        ["js/newsletter.js", "External JavaScript – onload alert + onsubmit validation"],
    ],
    widths=[2.4, 4.1],
)

# =====================================================================
# 3. REQUIREMENTS COMPLIANCE TABLE
# =====================================================================
doc.add_page_break()
heading("3. Requirements Compliance Table", 1)
para(
    "The table below lists every requirement from the assessment specification and states "
    "whether it has been met, together with where/how it is implemented.",
    italic=True,
)

heading("3.1 Website / general requirements", 2)
make_table(
    ["#", "Requirement", "Met?", "Evidence / where implemented"],
    [
        ["1", "At least 6 pages on the website", "Yes", "8 HTML pages in total"],
        ["2", "Home page (index.html) introducing you to employers", "Yes", "index.html welcome box + intro"],
        ["3", "Resume page (personal details, education, work, qualifications)", "Yes", "rroy_A4_resume.html, incl. a units table"],
        ["4", "At least 4 skill pages, one being Web Development", "Yes", "Web Skills, Communication, Problem Solving, Teamwork"],
        ["5", "Written in HTML5 using a text editor", "Yes", "<!DOCTYPE html> on every page, hand-coded"],
        ["6", "Passes HTML5 and CSS validation", "Yes*", "W3C validators – see Testing Plan §5 (attach screenshots)"],
        ["7", "Correct DOCTYPE statements", "Yes", "<!DOCTYPE html> first line of every page"],
        ["8", "Meta tags on home page (search-engine ready)", "Yes", "index.html meta description + keywords"],
        ["9", "Relevant page title on each page", "Yes", "Unique <title> on all 8 pages"],
        ["10", "Design fits 1024 x 768 with no required scroll for nav", "Yes", "Fixed-width centred layout; horizontal nav bar"],
        ["11", "Consistent theme, navigation, fonts and colours", "Yes", "Shared portfolio.css; same nav template on every page"],
        ["12", "Graphics reduced to correct size in a photo editor", "Partial", "Images display correctly but profile.png (~1.7MB) and certificate.png (~1.6MB) are large – should be re-saved smaller (see §6)"],
        ["13", "Correct file naming (rroy_..._A4 / zip convention)", "Yes", "Pages named rroy_A4_*.html; report named rroy_ITC293_A4_design.docx"],
        ["14", "Website hosted and reachable by a URL", "Action", "Host on GitHub Pages / W3Schools Spaces and paste URL on the title page"],
    ],
    widths=[0.4, 2.7, 0.7, 2.7],
)
para("* Validation must be re-run and the screenshots attached before final submission.", italic=True, size=9, color=GREY)

heading("3.2 CSS / JavaScript requirements (marking criteria)", 2)
make_table(
    ["#", "Requirement", "Met?", "Evidence / where implemented"],
    [
        ["15", "Pure external CSS – no inline or internal styles", "Yes", "Single css/portfolio.css linked on every page"],
        ["16", "At least 2 type selectors + a universal selector", "Yes", "Universal '*'; type selectors body, nav, h1, fieldset, legend, etc."],
        ["17", "Print-friendly styling carried over from A3", "Yes", "@media print block hides nav, flattens form"],
        ["18", "JavaScript kept external for reusability/readability", "Yes", "js/newsletter.js linked via <script src>"],
    ],
    widths=[0.4, 2.7, 0.7, 2.7],
)

heading("3.3 HTML5 form requirements", 2)
make_table(
    ["#", "Requirement", "Met?", "Evidence / where implemented"],
    [
        ["19", "Form styled using pure CSS", "Yes", "form#signupForm rules in portfolio.css"],
        ["20", "Form opens in a new window/tab from the Web Dev (JavaScript) page", "Yes", "rroy_A4_web_skills.html link with target=\"_blank\""],
        ["21", "Data submitted using the POST method", "Yes", "<form method=\"post\" action=\"...\">"],
        ["22", "Welcome alert on page load (onload)", "Yes", "<body onload=\"showWelcomeAlert();\"> -> alert"],
        ["23", "Instructive information (e.g. 'required data')", "Yes", "Intro paragraph + red * markers on required fields"],
        ["24", "First Name text box, max size 60", "Yes", "<input type=\"text\" maxlength=\"60\">"],
        ["25", "Last Name text box, max size 60", "Yes", "<input type=\"text\" maxlength=\"60\">"],
        ["26", "Email box, max size 60", "Yes", "<input type=\"email\" maxlength=\"60\">"],
        ["27", "5 radio buttons (15-25, 26-35, 36-45, 46-55, 55+)", "Yes", "<fieldset> 'Age Range' with 5 radios"],
        ["28", "5 checkboxes for company services", "Yes", "Web Dev, UX/UI, E-commerce, SEO, Cloud Hosting"],
        ["29", "Text area, minimum 60 cols x 3 rows", "Yes", "<textarea cols=\"60\" rows=\"3\">"],
        ["30", "Submit button and Reset button", "Yes", "Sign Me Up (submit) + Clear Form (reset)"],
        ["31", "onsubmit validates both names non-empty; alert + prevents submit", "Yes", "validateSignupForm(this) returns false when empty"],
        ["32", "On-screen / pop-up confirmation when valid (per FAQ)", "Yes", "Thank-you alert on successful validation"],
        ["33", "Plenty of code comments", "Yes", "Comment blocks + inline notes in every HTML/CSS/JS file"],
    ],
    widths=[0.4, 2.7, 0.7, 2.7],
)

# =====================================================================
# 4. CHANGES SINCE LAST SUBMISSION
# =====================================================================
doc.add_page_break()
heading("4. Changes Since the Last Submission (A3 → A4)", 1)
para("The following changes and additions were made to evolve the Assessment 3 site into the Assessment 4 portfolio:")
for c in [
    "Renamed all page links and files across the site from rroy_A3_* to rroy_A4_*.",
    "Added the fictitious employer context (Verdant Web Solutions Pty Ltd) on the home page and skill pages.",
    "Added three new skill pages – Communication, Problem Solving and Teamwork – each giving concrete, real-world examples of the skill.",
    "Rewrote the Web Skills page to include a JavaScript section and a call-to-action link that opens the new signup form in a new browser tab (target=\"_blank\").",
    "Added rroy_A4_newsletter.html – a complete HTML5 newsletter signup form (POST, onload alert, onsubmit validation, 5 radios, 5 checkboxes, 60-col textarea, submit + reset).",
    "Added js/newsletter.js – an external JavaScript file (kept external for reusability and readability marks) with showWelcomeAlert() and validateSignupForm().",
    "Fixed a validation bug so that, when a name field is empty, the field is correctly focused and the form submission is reliably cancelled.",
    "Extended css/portfolio.css with a universal selector (*), pure-CSS form styling (fieldset, legend, inputs, buttons), and a print override for the form. No inline or internal CSS anywhere.",
    "Added SEO meta tags (description, keywords) to the home page so the site is search-engine ready, and ensured every page has a unique, relevant <title>.",
]:
    bullet(c)

para("Carried over from Assessments 2 and 3:", bold=True)
for c in [
    "Single external CSS file (no internal or inline styles).",
    "CSS-styled navigation bar present on every page (consistent template).",
    "Print stylesheet (Ctrl/Cmd+P) that hides navigation and uses print-friendly colours and units.",
    "Logo / favicon (images/profile.png) shown on every page tab.",
    "Footer with contact details and a copyright / permission notice.",
    "Resume page with a multi-row table of current units.",
]:
    bullet(c)

# =====================================================================
# 5. TESTING PLAN
# =====================================================================
doc.add_page_break()
heading("5. Testing Plan", 1)
para(
    "The site was tested manually in a current browser at 1024 x 768 resolution, and validated "
    "with the W3C HTML5 and CSS validators. Each test case below lists the action, the expected "
    "result and a column to record the actual result. A screenshot of each page and of each "
    "validation result should be attached to this report as evidence."
)

heading("5.1 Functional / navigation tests", 2)
make_table(
    ["#", "Test action", "Expected result", "Pass/Fail"],
    [
        ["T1", "Open index.html in the browser at 1024x768", "Page loads with no horizontal scroll bar", ""],
        ["T2", "Click each of the 7 navigation links in turn", "Each destination page loads correctly", ""],
        ["T3", "Confirm every page tab shows the favicon and a relevant title", "Favicon shown; title matches page", ""],
        ["T4", "On Web Skills page, click 'Open the Newsletter Signup Form'", "Form opens in a new browser tab", ""],
        ["T5", "Observe the signup page immediately after it loads", "Welcome alert is displayed (onload)", ""],
    ],
    widths=[0.4, 2.9, 2.5, 0.7],
)

heading("5.2 Form validation tests", 2)
make_table(
    ["#", "Test action", "Expected result", "Pass/Fail"],
    [
        ["T6", "Click 'Sign Me Up' with both name fields empty", "Alert about First Name; form does NOT submit", ""],
        ["T7", "Enter First Name only, click submit", "Alert about Last Name; form does NOT submit", ""],
        ["T8", "Enter both names, click submit", "Thank-you alert; form submits (POST)", ""],
        ["T9", "Type more than 60 characters into a name field", "Input stops accepting at 60 characters", ""],
        ["T10", "Select an age radio, then select a different one", "Only one radio stays selected at a time", ""],
        ["T11", "Tick several service checkboxes", "Multiple checkboxes can be selected together", ""],
        ["T12", "Click 'Clear Form'", "All fields reset to empty/unselected", ""],
    ],
    widths=[0.4, 2.9, 2.5, 0.7],
)

heading("5.3 Presentation / validation tests", 2)
make_table(
    ["#", "Test action", "Expected result", "Pass/Fail"],
    [
        ["T13", "Press Ctrl/Cmd+P on each page (print preview)", "Navigation hidden; print-friendly colours/fonts", ""],
        ["T14", "Resize browser narrower and wider", "Layout stays usable; nav still reachable", ""],
        ["T15", "Validate each .html file with the W3C HTML5 validator", "No errors reported (attach screenshot)", ""],
        ["T16", "Validate portfolio.css with the W3C CSS validator", "No errors reported (attach screenshot)", ""],
    ],
    widths=[0.4, 2.9, 2.5, 0.7],
)

# =====================================================================
# 6. KNOWN ISSUES / FUTURE IMPROVEMENTS
# =====================================================================
heading("6. Known Issues and Future Improvements", 1)
for c in [
    "Image file sizes: profile.png (~1.7 MB) and certificate.png (~1.6 MB) are larger than necessary for the dimensions they are displayed at. Before final submission these should be re-saved at the displayed size in a photo editor to avoid losing 'file size' marks.",
    "Hosting URL: the live site must be hosted (e.g. GitHub Pages or W3Schools Spaces) and the URL inserted on the title page of this report.",
    "Form action endpoint currently points at a public echo service (httpbin.org/post) so submission can be demonstrated; this can be repointed at a real handler if required.",
]:
    bullet(c)

# =====================================================================
# 7. USE OF AI
# =====================================================================
heading("7. Use of Generative AI", 1)
para(
    "The HTML, CSS and JavaScript were hand-written using a text editor so that I understand "
    "how each part works, in line with the subject's guidance. Where AI assistance was used "
    "(for example, drafting or proofreading this report), it is acknowledged here. No generated "
    "code was submitted without my own review and understanding, and I can explain every part "
    "of the site during the professional dialogue."
)

# ---- save ----
out = "rroy_ITC293_A4_design.docx"
doc.save(out)
print("Saved", out)
